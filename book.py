from docx import Document

import time
import hashlib
import datetime

from sqlite3 import connect
from os import path, mkdir, listdir, remove

from PIL import ImageTk
from PIL import Image as IM


from sys import platform
from docx.shared import Pt

if platform == "linux" or platform == "linux2":
    usb = "/media/el/picture hard drive/"
    scan = "/media/el/VOLUME11/DCIM/100MEDIA"
    from Tkinter import *
    from tkMessageBox import askyesnocancel, showwarning
    from tkFileDialog import askopenfilename, asksaveasfilename, askdirectory
    KEYCODE = 22
    def Beep(*boop):
        print('\a')
        print('\a')
        print('\a')

        print("no beep? :(")

elif platform == "win32":
    usb = "E:\\"
    KEYCODE = 8
    scan = "F:\\DCIM\\100MEDIA"

    from tkinter import *
    from tkinter.messagebox import askyesnocancel, showwarning
    from tkinter.filedialog import askopenfilename, asksaveasfilename, askdirectory
    from winsound import Beep



no = "no image.png"
SQL = connect(path.join(usb, "data", "data.sqlite"))

#usb = "E:\\"
#usb = "/media/el/picture hard drive/"
hold = usb + "hold"
HASH = usb + "hashed"

class image(Frame):
    def __init__(self, root, id, x, y, th):
        Frame.__init__(self, root)
        self.id = id
        self.xy = x, y
        self.th = th
        self.grid(row=y, column=x)

        self.image_path = path.join(HASH, "no image.png")
        self.image = ImageTk.PhotoImage(IM.open(self.image_path).resize((100, 100), IM.ANTIALIAS))

        self.label = Label(self, image=self.image)
        self.label.pack()
        self.label.bind("<Button-1>", self.swap)

        self.text = Entry(self)
        self.text.bind("<Key>", self.lookup)
        self.text.bind("<Control-Up>", lambda a: self.set(self.xy[0], self.xy[1]-1))
        self.text.bind("<Control-Down>", lambda a: self.set(self.xy[0], self.xy[1]+1))
        self.text.bind("<Control-Left>", lambda a: self.set(self.xy[0]-1, self.xy[1]))
        self.text.bind("<Control-Right>", lambda a: self.set(self.xy[0]+1, self.xy[1]))
        #print(self.text.cget("background"))
        self.text.config({"background": "#ffffff"})
        self.text.bind("<FocusOut>", lambda x: self.text.config({"background": "#ffffff"}))
        self.text.pack()

        if id:
            self.found(self.id)
    def set(self, x, y):
        x = min(max(x, 0), 3)
        y = min(max(y, 0), 4)
        self.master.master.items[x][y].text.focus_set()

    def delim(self):
        self.id = None
        self.image_path = path.join(HASH, "no image.png")
        self.image = ImageTk.PhotoImage(IM.open(self.image_path).resize((100, 100), IM.ANTIALIAS))
        self.label.configure(image=self.image)

    def _newim(self, id):
        self.id = id
        self.image_path = path.join(HASH, id + ".JPG")
        self.image = ImageTk.PhotoImage(IM.open(self.image_path).resize((100, 100), IM.ANTIALIAS))
        self.label.configure(image=self.image)

    def lookup(self, key):
        #print(dir(key))
        t = self.text.get()
        if key.keycode != KEYCODE:
            t += key.char
        elif not t:
            self.delim()
        print("looking up", t, key.keycode, dir(key))
        global SQL
        c = SQL.cursor()
        c.execute("SELECT hash FROM images WHERE hash LIKE ?", (t + "%",))
        out = c.fetchall()
        out = set(out)
        print(out)
        if out:
            if len(out) == 1:
                self.text.config({"background": "#11aa11"})
                if key.keycode != KEYCODE:
                    Beep(1000, 200)
                    self.found(out.pop()[0])
            else:
                self.text.config({"background": "#aaaa11"})
        else:
            self.text.config({"background": "#aa1111"})
        c.close()

    def swap(self, butt):
        print(ANCHOR)
        scan_list = self.master.master.master.master.scan_list
        id = scan_list.get(ANCHOR)
        scan_list.delete(ANCHOR)
        if id:
            self.found(id)

    def found(self, id):

        self.text.delete(0, END)
        self.text.insert(0, id)
        self._newim(id)


class SlidePage(Frame):
    def __init__(self, root, page, data=[[("U", None) for a in range(5)] for a in range(4)]):
        Frame.__init__(self, root)
        self.page_name = page

        self.hold = Frame(self)
        self.hold.pack()

        self.label = Label(self, text=page)
        self.label.pack()

        self.items = [[image(self.hold, data[x][y][1], x, y, data[x][y][0]) for y in range(5)] for x in range(4)]

class NegPage(Frame):
    def __init__(self, root, page, data=[[("U", None) for a in range(7)] for a in range(6)]):
        Frame.__init__(self, root)
        self.page_name = page

        self.hold = Frame(self)
        self.hold.pack()

        self.label = Label(self, text=page)
        self.label.pack()

        self.items = [[image(self.hold, data[x][y][1], x, y, data[x][y][0]) for y in range(7)] for x in range(6)]


class PageHolder(Frame):
    def __init__(self):
        self.pages = []


class MenuBar(Menu):
    def __init__(self, root):

        root.config(menu=self)

        self.file = Menu(self, tearoff=0)
        self.file.add_command(label="Open", command=root.open)
        self.file.add_command(label="Save", command=root.save)
        self.file.add_command(label="Save as", command=root.saveas)
        self.file.add_command(label="Export docxs", command=root.export)
        self.file.add_separator()
        self.file.add_command(label="Exit", command=root.quit)
        self.add_cascade(label="File", menu=self)


class BookMaker(Tk):
    def __init__(self, name):
        Tk.__init__(self)
        self.ptype = SlidePage
        self.ptype = NegPage
        self.book_name = name

        self.menubar = Menu(self)
        self.config(menu=self.menubar)
        # create a pulldown menu, and add it to the menu bar
        self.filemenu = Menu(self.menubar, tearoff=0)
        self.filemenu.add_command(label="Open", command=self.open)
        self.filemenu.add_command(label="Save", command=self.save)
        self.filemenu.add_command(label="export docx", command=self.export)
        self.filemenu.add_command(label="Save as", command=self.saveas)
        self.filemenu.add_separator()
        self.filemenu.add_command(label="Exit", command=self.quit)
        self.menubar.add_cascade(label="File", menu=self.filemenu)

        self.side_hold = Frame(self)
        self.side_hold.pack(side=LEFT, fill=X, anchor="n")

        self.pages = []
        self.page_now = 0
        self.page_hold = Frame(self, height=2, width=2)
        self.page_hold.pack(side=LEFT, anchor="nw")
        self.scan_list = Listbox(self.side_hold)
        self.scan_list.pack(side=TOP, anchor="n")
        self.done = Button(self.side_hold, command=self.output, text="done")
        self.done.pack(side=RIGHT, anchor="s")
        self.new_page()
        self.idk = self.pages[0]
        self.pages[0].pack()

        self.pageflick = Frame(self.side_hold)
        self.pageflick.pack()
        self.next = Button(self.pageflick, command=self.next_page, text=">")
        self.new = Button(self.pageflick, command=self.new_page, text="O")
        self.back = Button(self.pageflick, command=self.back_page, text="<")
        self.next.pack(side=RIGHT)
        self.new.pack(side=RIGHT)
        self.back.pack(side=RIGHT)

        self.after(0, self.loop)

    def export(self):
        for i in self.pages:
            data = [[i.items[x][y].id for y in range(5)] for x in range(4)]
            make_book(i.page_name, self.book_name, data)

    def open(self):
        out = askyesnocancel("do you want to save?")
        book = askopenfilename()
        #book = "book1"
        print(out, book)
        if out:
            self.save()
            self.load_book2(book)
        elif out != None:
            self.load_book2(book)
        else:
            pass

    def save(self):
        self.save_book2()

    def saveas(self): # todo add
        pass

    def save_book(self):
        FILE = usb + "books"
        if not path.exists(path.join(FILE, self.book_name)):
            mkdir(path.join(FILE, self.book_name))
        for i in self.pages:
            with open(path.join(FILE, self.book_name, i.page_name + ".txt"), "w+") as file:
                for j in i.items:
                    for k in j:
                        if k.id:
                            file.write(k.id)
                        else:
                            file.write("_"*32)

    def save_book2(self):
        FILE = usb + "books"
        with open(path.join(FILE, self.book_name + ".book"), "w+") as file:
            file.write("001")
            if self.ptype is SlidePage:
                file.write("s")
            else:
                file.write("n")
            for i in self.pages:
                for j in i.items:
                    for k in j:
                        if k.id:
                            file.write(k.id)
                        else:
                            file.write("_"*32)
                        file.write(k.th)
                name = "_"*20
                name = i.page_name + name
                name = name[:20]
                file.write(name)

    def load_book2(self, book):
        pages = []
        #FILE = usb + "books"
        self.book_name = book
        self.title(self.book_name)
        self.page_now = 0
        with open(path.join(book), "r") as file:
            out = []
            if file.read(3) != "001":
                showwarning(title="file error", message="this file is not supported")
                return

            if file.read(1) == "s":
                X, Y = 4, 5
                pt = SlidePage
            else:
                X, Y = 6, 7
                pt = NegPage
            page = file.read(33 * X * Y)
            while page:

                if page:
                    out = []
                    for x in range(X):
                        out2 = []
                        for y in range(Y):
                            hash, page = page[:32], page[32:]
                            th, page = page[:1], page[1:]
                            if "_" in hash:
                                out2.append(("U", None))
                            else:
                                out2.append((th, hash))
                        out.append(out2)
                    name = file.read(20).replace("_", "")
                    pages.append(pt(self.page_hold, name, data=out))

                page = file.read(33 * X * Y)

        self.pages = pages
        self.turn_page()

    def load_book(self, book):
        pages = []
        FILE = usb + "books"
        self.book_name = book
        self.page_now = 0
        print(listdir(path.join(FILE, book)))
        for files in listdir(path.join(FILE, book)):

            if files.split(".")[-1] == "txt":
                with open(path.join(FILE, book, files), "r") as file:
                    out = []
                    for x in range(4):
                        out2 = []
                        for y in range(5):
                            hash = file.read(32)
                            #print(hash)
                            if "_" in hash:
                                out2.append(("U", None))
                            else:
                                out2.append(("U", hash))
                        out.append(out2)
                pages.append(self.ptype(self.page_hold, files.split(".")[0], data=out))

        self.pages = pages
        self.turn_page()

    def next_page(self):
        if len(self.pages)-1 > self.page_now:
            self.page_now += 1
        self.turn_page()

    def back_page(self):
        if self.page_now > 0:
            self.page_now -= 1
        self.turn_page()

    def output(self):
        data = self.pages[self.page_now]
        out = [[data.items[x][y].id for y in range(5)] for x in range(4)]
        make_book(data.page_name, self.book_name, out)

    def turn_page(self):
        for i in self.page_hold.winfo_children():
            i.pack_forget()
        self.pages[self.page_now].pack()

    def new_page(self):
        self.pages.append(self.ptype(self.page_hold, "page " + str(len(self.pages))))

    def press(self, x, y):
        print(x, y)

    def loop(self):
        t = time.time() + 1000
        state = 1
        while True:
            if t - time.time() > 0:
                if state == 1:
                    if path.exists(scan):
                        state = 2
                        print("------------------------------")
                        for i in get_scans():
                            self.scan_list.insert(0, i)
                            print(i)
                elif state == 2:
                    if not path.exists(scan):
                        state = 1

                t = time.time() + 1000

            self.update()

def md5(fname):
    hash_md5 = hashlib.md5()
    with open(fname, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

def add_db(folder, image):
    if not path.exists(path.join(folder, "scan.txt")):
        with open(path.join(folder, "scan.txt"), "w+") as file:
            pass
    with open(path.join(folder, "scan.txt"), "r") as file:
        items = [i for i in file.read().splitlines()]
    if image in items:
        return False

    file = path.join(folder, image)
    HASHED = usb + "hashed"

    h = md5(file)
    with open(file, "rb") as f1:
        with open(path.join(HASHED, h + "." + image.split(".")[-1]), "wb+") as f2:
            f2.write(f1.read())
    d = path.getmtime(file)
    d = datetime.datetime.utcfromtimestamp(d)
    c = SQL.cursor()
    c.execute("INSERT INTO images (hash, ex, date) VALUES (?,?,?)", (h, image.split(".")[-1], d,))
    SQL.commit()
    c.close()

    with open(path.join(folder, "scan.txt"), "w+") as file:
        for i in items:
            file.write(i)
            file.write("\n")
    return True

def make_book(name, book, ids):
    doc2 = Document()

    section = doc2.sections[0]
    section.left_margin = int(section.left_margin / 2)
    section.right_margin = int(section.right_margin / 2)
    h = section.page_width - (section.left_margin + section.right_margin)
    h /= 4

    section.top_margin = int(section.top_margin / 3)
    section.bottom_margin = int(section.top_margin / 3)

    t = doc2.sections[0].header.paragraphs[0]
    #t.text = name
    t.alignment = 1
    t.bold = True
    t.add_run(name).font.size = Pt(40)

    for y in range(5):
        par = doc2.add_paragraph("")
        par.paragraph_format.line_spacing = 0
        par.paragraph_format.line_spacing_rule = 3
        par.paragraph_format.space_after = 0
        par.paragraph_format.space_before = 0
        for x in range(4):
            if ids[x][y]:
                par.add_run().add_picture(path.join(HASH, ids[x][y] + ".JPG"), width=h,
                                      height=h/1.5)
            else:
                par.add_run().add_picture(path.join(HASH, "no image.png"), width=h,
                                          height=h/1.5)
        table = doc2.add_table(rows=1, cols=4)
        # table.style = 'Table Grid'

        cell = table.rows[0].cells
        for x in range(4):
            p = cell[x].paragraphs[0]
            p.alignment = 1
            if ids[x][y]:
                p.add_run(ids[x][y][:9]).bold = True
    FILE = usb + "books"
    book = path.join(FILE, book)
    if not path.exists(book):
        mkdir(book)

    doc2.save(path.join(book, name + '.docx'))

def get_scans():


    save = usb + "pictures"
    tt = str(time.time())
    if path.exists(scan):
        # print("path good")
        for f in listdir(scan):
            out1 = ""
            out1 = ""  # askstring("rename", str(f))
            if not out1:
                out = str(f)
            else:
                out = out1 + f.split(".")[-1]
            l = path.join(save, tt)
            if not path.exists(l):
                mkdir(l)
            h = md5(path.join(scan, f))
            dupe_move(path.join(scan, f), path.join(hold, h + "." + f.split(".")[-1]))
            move(path.join(scan, f), path.join(l, out))
            add_db(l, out)
            yield h
            #add(md5(path.join(l, out)), f, out1)
            # print(md5(os.path.join(l, out + f.split(".")[-1])), f, out)
    else:
        print("path bad")


def move(a, b):
    with open(a, "rb") as f1:
        with open(b, "wb+") as f2:
            f2.write(f1.read())
    remove(a)

def dupe_move(a, b):
    with open(a, "rb") as f1:
        with open(b, "wb+") as f2:
            f2.write(f1.read())




BookMaker("negtestbook").mainloop()



#make_book("test 1", "test" , [["0a5d53976216fcd2db71decf8bea95d1" for y in range(5)] for x in range(4)])